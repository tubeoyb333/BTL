import tkinter as tk
from tkinter import ttk, simpledialog, messagebox,filedialog
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import matplotlib.pyplot as plt
import tkinter.filedialog
import sympy as sp
import numpy as np
import pandas as pd
import os
from docx import Document

class MathExpression:
    def __init__(self, expression_str):
        self.expression_str = expression_str
        self.x = sp.symbols('x')
        try:
            self.expression = sp.sympify(expression_str)
        except Exception as e:
            raise ValueError(f"Lỗi: {str(e)}")
class MathAppGUI:
    def __init__(self, master):
        self.master = master
        self.master.title("Math Functions GUI")

        self.math_expression = None
        self.functions_from_excel = []
        self.extreme_points = []
        self.results_history = []
        
        self.create_widgets()

    def create_widgets(self):
        bigger_font = ("Helvetica", 15)
        ttk.Label(self.master, text="Nhập hàm số (sử dụng 'x' là biến):",font=bigger_font).grid(row=0, column=0, pady=10, padx=10, sticky=tk.W)

        self.entry_function = ttk.Entry(self.master, width=50)
        self.entry_function.grid(row=0, column=1, pady=10, padx=10, columnspan=2)

        ttk.Button(self.master, text="Nhập Dữ Liệu từ Excel", command=self.load_data_from_excel).grid(row=0, column=3, pady=10, padx=10, sticky=tk.W)

        functions = {
            "Tính Đạo Hàm": self.calculate_derivative,
            "Tính Nguyên Hàm": self.calculate_integral,
            "Giải Phương Trình": self.solve_equation,
            "Tính Diện Tích Hàm Số và Đường Thẳng": self.calculate_area,
            "Tìm Điểm Cực Trị": self.find_extreme_points,
            "Tính Diện Tích Giữa Hai Hàm": self.calculate_area_between_functions,
            "Kiểm Tra Tính Liên Tục": self.check_continuity,
            "Lưu Kết Quả": self.save_result
        }
        row_index = 1
        for text, command in functions.items():
            ttk.Button(self.master, text=text, command=command, padding=(5, 15), width=50).grid(row=row_index, column=0, pady=0, padx=10, sticky=tk.W)
            row_index += 1

        self.current_function_label = ttk.Label(self.master, text="Hàm số đang được sửa đổi:",font=bigger_font)
        self.current_function_label.grid(row=1, column=1, pady=10, padx=10, columnspan=2, sticky=tk.W)

        self.result_label = ttk.Label(self.master, text="", wraplength=400, background="lightgray", font=("Arial", 12))
        self.result_label.grid(row=10, column=0, pady=10, padx=10, columnspan=3, rowspan=2, sticky=tk.W + tk.E)

        self.figure = Figure(figsize=(5, 4), dpi=100)
        self.plot = self.figure.add_subplot(1, 1, 1)
        self.canvas = FigureCanvasTkAgg(self.figure, master=self.master)
        self.canvas.get_tk_widget().grid(row=2, column=2, rowspan=8, pady=50, padx=50, columnspan=2)
        self.entry_function.bind("<KeyRelease>", self.update_plot)
        
    def update_plot(self, event):
        expression_str = self.entry_function.get()
        self.math_expression = MathExpression(expression_str)
        x_vals = np.linspace(-10, 10, 1000)
        y_vals = [self.math_expression.expression.subs(self.math_expression.x, val) for val in x_vals]
        self.plot.clear()
        self.plot.plot(x_vals, y_vals, label='Hàm số')

        for point in self.extreme_points:
            x_val = point['point']
            y_val = self.math_expression.expression.subs(self.math_expression.x, x_val)
            self.plot.scatter([x_val], [y_val], color='red', label=f'{point["type"]} tại x={x_val}')

        self.plot.set_title("Đồ thị hàm số")
        self.plot.set_xlabel("x")
        self.plot.set_ylabel("f(x)")
        self.plot.legend()
        self.canvas.draw()
        
        self.current_function_label.config(text=f"Hàm số đang được sửa đổi: {expression_str}")
    def load_data_from_excel(self):
        file_path = filedialog.askopenfilename()
        if file_path:
            try:
                data_frame = pd.read_excel(file_path)

                if 'function' in data_frame.columns and not data_frame['function'].empty:

                    self.functions_from_excel = data_frame['function'].tolist()

                    function_combobox = ttk.Combobox(self.master, values=self.functions_from_excel, state="readonly")
                    function_combobox.grid(row=0, column=4, pady=10, padx=10)
                    function_combobox.set("Chọn hàm số")

                    function_combobox.bind("<<ComboboxSelected>>", self.selected_function_from_combobox)
                data_preview = data_frame.head()
                self.show_result(f"Nội dung từ tệp tin:\n{data_preview}")
            except Exception as e:
                self.show_error(f"Lỗi khi đọc tệp tin Excel: {str(e)}")

    def selected_function_from_combobox(self, event):
        selected_function = event.widget.get()
        if selected_function:
            self.entry_function.delete(0, tk.END)
            self.entry_function.insert(0, selected_function)
            self.update_plot(None)
            self.show_result(f"Đã chọn hàm số từ danh sách: {selected_function}")

    def validate_variable(self, expression_str):
        try:
            expression = sp.sympify(expression_str)
            variables = expression.free_symbols
            if len(variables) == 1 and sp.symbols('x') in variables:
                return True
            else:
                self.show_error("Biến phải là 'x'.")
                return False
        except Exception:
            self.show_error("Biểu thức không hợp lệ.")
            return False
    
    def calculate_derivative(self):
        expression_str = self.entry_function.get()
        if not self.validate_variable(expression_str):
            return
        try:
            math_expression = MathExpression(expression_str)
            x = math_expression.x
            expression = math_expression.expression
            n = self.get_single_value("Bậc của đạo hàm cần tính:")
            derivate = sp.diff(expression, x, n)
            self.show_result(f"Đạo hàm bậc {n}: {derivate}")
        except ValueError:
            self.show_error("Vui lòng nhập một số nguyên làm bậc của đạo hàm.")
        
    def calculate_integral(self):
        expression_str = self.entry_function.get()
        if not self.validate_variable(expression_str):
            return
        math_expression = MathExpression(expression_str)
        x = math_expression.x
        expression = math_expression.expression
        original = sp.integrate(expression, x)
        self.show_result(f"Nguyên hàm của hàm số: {original}")

    def solve_equation(self):
        expression_str = self.entry_function.get()
        if not self.validate_variable(expression_str):
            return
        math_expression = MathExpression(expression_str)
        x = math_expression.x
        expression = math_expression.expression
        solutions = sp.solve(expression, x)
        self.show_result(f"Nghiệm của phương trình: {solutions}")

    def calculate_area(self):
        expression_str = self.entry_function.get()
        if not self.validate_variable(expression_str):
            return
        range_values = self.get_values("Nhập khoảng giá trị x (bắt đầu và kết thúc, cách nhau bởi dấu cách):")
        if not all(val.isnumeric() for val in range_values):
            self.show_error("Vui lòng nhập các giá trị số cho khoảng.")
            return
        range_start, range_end = float(range_values[0]), float(range_values[1])
        if range_end <= range_start:
            self.show_error("Số kết thúc phải lớn hơn số trước đó.")
            return None
        math_expression = MathExpression(expression_str)
        x = math_expression.x
        expression = math_expression.expression
        area = sp.integrate(expression, (x, range_start, range_end))
        self.show_result(f"Diện tích hình phẳng giới hạn bởi hàm số và 2 đường thẳng: {area}")

    def find_extreme_points(self):
        expression_str = self.entry_function.get()
        if not self.validate_variable(expression_str):
            return
        math_expression = MathExpression(expression_str)
        x = math_expression.x
        expression = math_expression.expression
        critical_points = sp.solveset(sp.diff(expression, x), x, domain=sp.S.Reals)
        self.extreme_points = []
        has_extreme_point = False
        for point in critical_points:
            second_derivative = sp.diff(expression, x, 2).subs(x, point)
            if second_derivative > 0:
                self.extreme_points.append({"point": point, "type": "Cực tiểu"})
                has_extreme_point = True
            elif second_derivative < 0:
                self.extreme_points.append({"point": point, "type": "Cực đại"})
                has_extreme_point = True
        result_message = "Các điểm cực trị của hàm số:\n"
        if has_extreme_point:
            for point in self.extreme_points:
                result_message += f"{point['type']} tại x = {point['point']}\n"
        else:
            result_message = "Không có điểm cực trị trong phạm vi xác định."
        self.show_result(result_message)
        self.update_plot(None)
    
    def calculate_area_between_functions(self):
        expression_str = self.entry_function.get()
        x, expression_1 = self.math_expression.x, self.math_expression.expression
        if not self.validate_variable(expression_str):
            return
        expression_str_2 = self.get_single_value("Nhập hàm số thứ 2:")
        x_2, expression_2 = self.math_expression.x, sp.sympify(expression_str_2)
        if not self.validate_variable(expression_str_2):
            return
        intersections = sp.solve(sp.Eq(expression_1, expression_2), x)
        if len(intersections) != 2:
            self.show_error("Hai hàm số không giao nhau đúng hai điểm.")
            return
        intersections.sort()
        area = sp.integrate(abs(expression_1 - expression_2), (x, intersections[0], intersections[1]))
        self.show_result(f"Diện tích giới hạn bởi hai hàm số: {area}")

    def check_continuity(self):
        expression_str = self.entry_function.get()
        if not self.validate_variable(expression_str):
            return
        math_expression = MathExpression(expression_str)
        x = math_expression.x
        expression = math_expression.expression
        a = self.get_single_value("Nhập giá trị cần kiểm tra tính liên tục tại x = :")
        limit_left = sp.limit(expression, x, a, dir='-')
        limit_right = sp.limit(expression, x, a, dir='+')

        if limit_left == limit_right:
            self.show_result(f"Hàm số liên tục tại x = {a}")
        else:
            self.show_result(f"Hàm số không liên tục tại x = {a}")
 
    def show_result(self, message):

        self.result_label.config(text=message)
        self.results_history.append(message)

    def show_error(self, message):
        messagebox.showerror("Lỗi", message)

    def get_values(self, message):
        return self.get_single_value(message).split()

    def get_single_value(self, message):
        return simpledialog.askstring("Input", message)

    def save_result(self):
        if not self.results_history:
            self.show_error("Không có kết quả để lưu.")
            return
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx"), ("All files", "*.*")])
        if file_path:
            try:
                doc = Document()
                doc.add_heading("Kết Quả Tính Toán", level=1)
                initial_function = self.current_function_label.cget("text")
                doc.add_heading("Hàm số đang được sửa đổi:", level=2)
                doc.add_paragraph(initial_function)
                for result in self.results_history:
                    doc.add_paragraph(result)
                doc.save(file_path)
                self.show_result(f"Kết quả đã được lưu tại: {file_path}")
                graph_file_path = os.path.splitext(file_path)[0] + "_graph.png"
                self.figure.savefig(graph_file_path)
                self.show_result(f"Đồ thị hàm số đã được lưu tại: {graph_file_path}")
            except Exception as e:
                self.show_error(f"Lỗi khi lưu kết quả: {str(e)}")
def main():
    root = tk.Tk()
    app = MathAppGUI(root)
    root.mainloop()
if __name__ == "__main__":
    main()
